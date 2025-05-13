from django.http import HttpResponse
from docx import Document
from datetime import datetime, timedelta
from rest_framework import generics, permissions
from .serializers import IncomeSerializer, ExpenseSerializer
from django.shortcuts import render, redirect
from .forms import IncomeForm, ExpenseForm
from django.contrib.auth.decorators import login_required
from finance.models import Income, Expense
from django.db.models import Sum

class IncomeListCreateView(generics.ListCreateAPIView):
    serializer_class = IncomeSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return Income.objects.filter(user=self.request.user).order_by('-date')

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


class ExpenseListCreateView(generics.ListCreateAPIView):
    serializer_class = ExpenseSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        return Expense.objects.filter(user=self.request.user).order_by('-date')

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


def export_report(request, period):
    user = request.user
    if not user.is_authenticated:
        return HttpResponse("Unauthorized", status=401)

    today = datetime.today().date()

    if period == "weekly":
        start_date = today - timedelta(days=7)
    elif period == "monthly":
        start_date = today.replace(day=1)
    elif period == "yearly":
        start_date = today.replace(month=1, day=1)
    else:
        return HttpResponse("Invalid period", status=400)

    incomes = Income.objects.filter(user=user, date__gte=start_date)
    expenses = Expense.objects.filter(user=user, date__gte=start_date)

    doc = Document()
    doc.add_heading(f"{period.capitalize()} Report", 0)

    doc.add_heading("Incomes", level=1)
    for inc in incomes:
        doc.add_paragraph(f"{inc.date} - {inc.amount} ({inc.income_type}) [{inc.payment_method}]")

    doc.add_heading("Expenses", level=1)
    for exp in expenses:
        doc.add_paragraph(f"{exp.date} - {exp.amount} ({exp.expense_type}) [{exp.payment_method}]")

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename={period}_report.docx'
    doc.save(response)
    return response




# @login_required
# def home_view(request):
#     incomes = Income.objects.filter(user=request.user).order_by('-date')[:5]
#     expenses = Expense.objects.filter(user=request.user).order_by('-date')[:5]
#
#     context = {
#         'incomes': incomes,
#         'expenses': expenses
#     }
#     return render(request, 'index.html', context)



@login_required
def home_view(request):
    user = request.user
    incomes = Income.objects.filter(user=user)
    expenses = Expense.objects.filter(user=user)

    total_income = incomes.aggregate(total=Sum('amount'))['total'] or 0
    total_expense = expenses.aggregate(total=Sum('amount'))['total'] or 0
    balance = total_income - total_expense

    context = {
        'incomes': incomes,
        'expenses': expenses,
        'balance': balance,
    }
    return render(request, 'index.html', context)



@login_required
def income_add_view(request):
    if request.method == 'POST':
        form = IncomeForm(request.POST)
        if form.is_valid():
            income = form.save(commit=False)
            income.user = request.user
            income.save()
            return redirect('home')
    else:
        form = IncomeForm()
    return render(request, 'income_add.html', {'form': form})




@login_required
def expense_add_view(request):
    if request.method == 'POST':
        form = ExpenseForm(request.POST)
        if form.is_valid():
            expense = form.save(commit=False)
            expense.user = request.user
            expense.save()
            return redirect('/')
    else:
        form = ExpenseForm()
    return render(request, 'expense_add.html', {'form': form})
